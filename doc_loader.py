import os
import boto3
import re
from docx import Document
import PyPDF2
import fitz
from PIL import Image
import json
from pathlib import Path




TEXT_SAVE_DIR= Path("static/text")
IMAGE_SAVE_DIR=Path("static/images")
TEXT_SAVE_DIR.mkdir(parents=True ,exist_ok=True)
IMAGE_SAVE_DIR.mkdir(parents=True ,exist_ok=True)



def preprocess_text(text):
    text = text.lower()
    # Allow spaces in addition to the existing allowed characters
    allowed = r"a-z0-9\.\,\!\?\s"
    text = re.sub(rf"[^\n{allowed}]", "", text)
    # Replace multiple spaces/tabs with a single space
    text = re.sub(r'[\t ]+', ' ', text)
    # Clean up lines while preserving paragraphs
    text = '\n'.join([line.strip() for line in text.splitlines() if line.strip()])
    return text

def load_docx(file_path):
    doc=Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def load_pdf(file_path):
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        return "\n".join([page.extract_text() or '' for page in pdf_reader.pages])

def load_txt(file_path):
    with open(file_path,'r',encoding='utf-8') as f:
        return f.read()


def extract_images_from_docx(file_path):
    """Extracts Image From The DOC files"""
    doc=Document(file_path)
    rels=doc.part._rels
    count=0
    saved_path=[]

    for rel in rels:
        rel=rels[rel]
        if "image" in rel.target_ref:
            count+=1
            image_data=rel.target_part.blob
            output_path=os.path.join(IMAGE_SAVE_DIR,f"image_{count}.png")

            # breakpoint()
            with open(output_path,'wb') as f:
                f.write(image_data)
            saved_path.append(output_path)
    return saved_path

def extract_images_from_pdf(file_path):
    """Extracts Image From The pdf"""
    doc=fitz.open(file_path)
    img_count=0
    saved_paths=[]

    for page_index in range(len(doc)):
        page=doc[page_index]
        images=page.get_images(full=True)
        for img_index,img in enumerate(images):
            xref=img[0]
            base_image=doc.extract_image(xref)
            img_bytes=base_image["image"]
            ext= base_image["ext"]
            img_count+=1
            image_path=os.path.join(IMAGE_SAVE_DIR,f"pdf_image_{page_index+1}_{img_index+1}.{ext}")
            with open(image_path,'wb') as f:
                f.write(img_bytes)
                saved_paths.append(image_path)

    return saved_paths



"""def upload_to_s3(text,bucket_name,s3_key,image_paths=None):
    Upload to S3 bucket
    s3=boto3.client('s3')
    s3.put_object(Body=text.encode('utf-8'),Bucket=bucket_name,Key=s3_key)
    print(f"Uploaded file to S3 bucket {bucket_name} with key {s3_key}" )
    if image_paths:
        for local_path in image_paths:
            file_name=os.path.basename(local_path)
            image_key=f"{os.splitext(s3_key)[0]}_images/{file_name}"
            with open(local_path,"rb") as img_file:
                s3.upload_fileobj(img_file,bucket_name,image_key)
            print(f"Uploaded text to s3://{bucket_name}/{s3_key}")
"""

def load_and_process_doc(file_path):
    ext=os.path.splitext(file_path)[1].lower()
    if ext==".pdf":
        raw_text=load_pdf(file_path)
        images=extract_images_from_pdf(file_path)
    elif ext==".docx":
        raw_text=load_docx(file_path)
        images=extract_images_from_docx(file_path)
    elif ext==".txt":
        raw_text=load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    clean_text=preprocess_text(raw_text)
    return clean_text,images


if __name__ == "__main__":
    file_path="test.docx"


    text, image_paths = load_and_process_doc(file_path)
    print("Preview of processed text:\n", text[:300])
    filename=os.path.splitext(os.path.basename(file_path))[0]
    text_output_path=os.path.join(TEXT_SAVE_DIR,f"{filename}.json")
    with open(text_output_path,'w',encoding='utf-8') as f:
        json.dump({"text":text,"image_paths":image_paths},f,indent=2)
    print(f"✅ Saved cleaned text to: {text_output_path}")
    print(f"✅ Extracted and saved {len(image_paths)} image(s) to: {IMAGE_SAVE_DIR}")
